#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_mtbf_docx.py  —  MTBF 计算书 Word 文档生成器
用法：
    python gen_mtbf_docx.py '<JSON_CONFIG_STRING>'
或在代码中：
    from gen_mtbf_docx import run
    run(config_dict)

CONFIG 字段说明见文件末尾 DEFAULT_CONFIG。
"""

import sys
import json
import io
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# MIL-HDBK-217F GF 环境 40°C 失效率数据（×10⁻⁶/h）
# ─────────────────────────────────────────────────────────────────────────────
LG = {
    "fuse":     0.0100,   # Section 12 低压熔断器
    "mcb":      0.0100,   # Section 11 微型断路器/直流开关
    "mov":      0.0030,   # Section 26 MOV/TVS 防雷器
    "busbar":   0.0010,   # Section 17 铜排/母排（工程估算）
    "terminal": 0.0027,   # Section 15 接插件/端子排
    "hall":     0.0100,   # Section 5  霍尔传感器（线性IC）
    "relay":    0.0400,   # Section 13 小型密封继电器
    "led":      0.0100,   # Section 6  LED 指示灯
}

# 成套组件 MTBF（h）
SUBUNIT_MTBF = {
    "hmi_7in":   300_000,
    "hmi_10in":  250_000,
    "ctrl_board": 800_000,
    "din_psu":  2_000_000,
    "imd":        600_000,
}

# 器件质量等级 → πQ
GRADE_PQ = {
    "国产工业级": 2.0,
    "进口工业级": 1.5,
    "进口品牌":   1.5,
    "imported":   1.5,
    "军品级":     1.0,
    "mil":        1.0,
}

# 特殊器件 πQ（不随 component_grade 变化）
PQ_BUSBAR = 1.0    # 纯导体，无电子失效机制
PQ_LED    = 1.5    # 工业级 LED

WIRING_SS1 = 0.200   # 配电子系统内部线缆估算值
WIRING_SS2 = 0.300   # 监控子系统内部线缆估算值


def _subunit_lam(mtbf_h):
    return round(1e6 / mtbf_h, 4)


# ─────────────────────────────────────────────────────────────────────────────
# 计算引擎
# ─────────────────────────────────────────────────────────────────────────────
def calculate(cfg):
    inp     = cfg["input_config"]
    out     = cfg["output_config"]
    lp      = cfg["lightning_protection"]
    mon     = cfg["monitoring"]
    req_h   = cfg.get("mtbf_requirement_h", 100_000)
    pq_i    = GRADE_PQ.get(cfg.get("component_grade", "国产工业级"), 2.0)

    # ── 子系统一：配电 ────────────────────────────────────────────────────────
    n_in  = inp["count"]
    n_lp  = lp["count"]

    input_type = inp.get("type", "熔断器")
    input_lg   = LG["fuse"] if "熔断器" in input_type else LG["mcb"]

    # 输出支路：支持单组（旧格式）和多组（groups）
    if "groups" in out:
        out_groups = out["groups"]
    else:
        out_groups = [{
            "count":  out["total_count"],
            "rating": out["rating_per_circuit"],
            "poles":  out["poles"],
            "type":   out.get("type", "熔断器"),
        }]
    n_out_total = sum(g["count"] for g in out_groups)

    # 构建 SS1 行（序号动态）
    ss1_items = []
    seq = 1

    # 输入保护
    inp_single = round(input_lg * pq_i, 4)
    inp_total  = round(n_in * inp_single, 4)
    ss1_items.append({
        "seq": seq, "seq_auto": True,
        "desc": f'输入{input_type}（{inp["rating"]}）',
        "n": n_in, "lg": input_lg, "pq": pq_i,
        "single": inp_single, "total": inp_total,
        "ref": "Sec.12/11",
    })
    seq += 1

    # 输出保护（每组一行）
    for grp in out_groups:
        g_type   = grp.get("type", "熔断器")
        g_lg     = LG["fuse"] if "熔断器" in g_type else LG["mcb"]
        g_single = round(g_lg * pq_i, 4)
        g_total  = round(grp["count"] * g_single, 4)
        ss1_items.append({
            "seq": seq, "seq_auto": True,
            "desc": f'输出{g_type}（{grp["rating"]}/{grp["poles"]}P）',
            "n": grp["count"], "lg": g_lg, "pq": pq_i,
            "single": g_single, "total": g_total,
            "ref": "Sec.12/11",
        })
        seq += 1

    # 防雷器
    lp_single = round(LG["mov"] * pq_i, 4)
    lp_total  = round(n_lp * lp_single, 4)
    ss1_items.append({
        "seq": seq, "desc": f'直流防雷器（{lp["rating_kA"]}kA，{lp.get("waveform","8/20μs")}）',
        "n": n_lp, "lg": LG["mov"], "pq": pq_i,
        "single": lp_single, "total": lp_total, "ref": "Sec.26",
    })
    seq += 1

    # 铜排
    bus_single = round(LG["busbar"] * PQ_BUSBAR, 4)
    bus_total  = round(4 * bus_single, 4)
    ss1_items.append({
        "seq": seq, "desc": "铜排/汇流母排",
        "n": 4, "lg": LG["busbar"], "pq": PQ_BUSBAR,
        "single": bus_single, "total": bus_total, "ref": "Sec.17",
    })
    seq += 1

    # 输出端子
    n_out_term  = n_out_total
    ot_single   = round(LG["terminal"] * pq_i, 4)
    ot_total    = round(n_out_term * ot_single, 4)
    ss1_items.append({
        "seq": seq, "desc": "输出接线端子/铜鼻子",
        "n": n_out_term, "lg": LG["terminal"], "pq": pq_i,
        "single": ot_single, "total": ot_total, "ref": "Sec.15",
    })
    seq += 1

    # 输入端子
    n_in_term  = n_in * 2
    it_single  = round(LG["terminal"] * pq_i, 4)
    it_total   = round(n_in_term * it_single, 4)
    ss1_items.append({
        "seq": seq, "desc": "输入电缆夹/接线端子",
        "n": n_in_term, "lg": LG["terminal"], "pq": pq_i,
        "single": it_single, "total": it_total, "ref": "Sec.15",
    })
    seq += 1

    # 内部线缆（固定估算）
    ss1_items.append({
        "seq": seq, "desc": "内部连接线缆（含绝缘）",
        "n": None, "lg": None, "pq": None,
        "single": None, "total": WIRING_SS1, "ref": "经验值",
    })

    lambda1 = round(sum(item["total"] for item in ss1_items), 4)

    # 构建 SS1 表行（供文档使用）
    ss1_rows = []
    for item in ss1_items:
        if item["n"] is None:
            ss1_rows.append((str(item["seq"]), item["desc"], "1", "—", "—", "—", f"{item['total']:.4f}"))
        else:
            ss1_rows.append((
                str(item["seq"]), item["desc"],
                str(item["n"]),
                f"{item['lg']:.4f}",
                f"{item['pq']}",
                f"{item['single']:.4f}",
                f"{item['total']:.4f}",
            ))
    ss1_rows.append(("**", "**子系统一（配电子系统）失效率合计 λ1**", "", "", "", "", f"**{lambda1:.4f}**"))

    # ── 子系统二：监控 ────────────────────────────────────────────────────────
    hmi_key  = "hmi_10in" if mon.get("hmi_size", "7寸").startswith("10") else "hmi_7in"
    hmi_lam  = _subunit_lam(SUBUNIT_MTBF[hmi_key])
    ctrl_lam = _subunit_lam(SUBUNIT_MTBF["ctrl_board"])
    psu_lam  = _subunit_lam(SUBUNIT_MTBF["din_psu"])
    has_imd  = bool(mon.get("has_insulation_monitor"))
    imd_lam  = _subunit_lam(SUBUNIT_MTBF["imd"]) if has_imd else 0.0

    n_hall  = mon.get("hall_sensor_count", n_in * 2)
    n_relay = mon.get("relay_count", 4)
    n_led   = mon.get("led_count", 10)

    hall_total  = round(n_hall  * LG["hall"]  * pq_i, 4)
    relay_total = round(n_relay * LG["relay"] * pq_i, 4)
    led_total   = round(n_led   * LG["led"]   * PQ_LED, 4)

    ss2_items = [
        {"seq": 1, "desc": f'{mon.get("hmi_size","7寸")}工业触摸屏显示单元',
         "n": "1", "ref": f'MTBF≥{SUBUNIT_MTBF[hmi_key]//1000}kh（制造商声明）', "pq": "—", "total": hmi_lam},
        {"seq": 2, "desc": "监控主控板（MCU+通信+I/O）",
         "n": "1", "ref": "MTBF≥800kh（设计估算）", "pq": "—", "total": ctrl_lam},
        {"seq": 3, "desc": "导轨式开关电源（监控供电）",
         "n": "1", "ref": "MTBF≥2,000kh（Telcordia）", "pq": "—", "total": psu_lam},
    ]
    ss2_seq = 4
    if has_imd:
        ss2_items.append({
            "seq": ss2_seq, "desc": "直流分路绝缘监测模块（IMD）",
            "n": "1", "ref": "MTBF≥600kh（制造商声明）", "pq": "—", "total": imd_lam,
        })
        ss2_seq += 1

    ss2_items += [
        {"seq": ss2_seq,   "desc": "霍尔电流传感器（主路监测）",
         "n": str(n_hall),  "ref": f"λ_G={LG['hall']:.4f}", "pq": f"{pq_i}", "total": hall_total},
        {"seq": ss2_seq+1, "desc": "告警继电器（小型密封）",
         "n": str(n_relay), "ref": f"λ_G={LG['relay']:.4f}", "pq": f"{pq_i}", "total": relay_total},
        {"seq": ss2_seq+2, "desc": "LED 状态指示灯",
         "n": str(n_led),   "ref": f"λ_G={LG['led']:.4f}",  "pq": f"{PQ_LED}", "total": led_total},
        {"seq": ss2_seq+3, "desc": "监控内部连接线缆及PCB",
         "n": "—", "ref": "经验值", "pq": "—", "total": WIRING_SS2},
    ]

    # 全局序号（接着 SS1 末尾）
    ss1_count = len(ss1_items)
    ss2_rows = []
    for i, item in enumerate(ss2_items):
        global_seq = ss1_count + i + 1
        ss2_rows.append((
            str(global_seq), item["desc"],
            item["n"], item["ref"], item["pq"],
            f"{item['total']:.4f}",
        ))

    lambda2 = round(sum(item["total"] for item in ss2_items), 4)
    ss2_rows.append(("**", "**子系统二（监控子系统）失效率合计 λ2**", "", "", "", f"**{lambda2:.4f}**"))

    lambda_total = round(lambda1 + lambda2, 4)
    mtbf_h       = round(1e6 / lambda_total)
    margin_h     = mtbf_h - req_h

    return {
        "ss1_rows":      ss1_rows,
        "ss2_rows":      ss2_rows,
        "lambda1":       lambda1,
        "lambda2":       lambda2,
        "lambda_total":  lambda_total,
        "mtbf_h":        mtbf_h,
        "mtbf_wan":      round(mtbf_h / 10000, 2),
        "mtbf_years":    round(mtbf_h / 8760,  1),
        "pct_ss1":       round(lambda1 / lambda_total * 100, 1),
        "pct_ss2":       round(lambda2 / lambda_total * 100, 1),
        "req_h":         req_h,
        "margin_h":      margin_h,
        "margin_pct":    round(margin_h / req_h * 100, 1),
        "pass":          margin_h >= 0,
        "pq_i":          pq_i,
        "has_imd":       has_imd,
        "out_groups":    out_groups,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Word 文档辅助函数
# ─────────────────────────────────────────────────────────────────────────────
def _font(run, bold=False, size=11, color=None):
    run.font.name = '宋体'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _para(doc, text, bold=False, size=11,
          align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    _font(r, bold=bold, size=size, color=color)
    return p


def _heading(doc, text, level=1):
    sizes   = {1: 16, 2: 13, 3: 12}
    befores = {1: 12, 2: 8,  3: 6}
    return _para(doc, text, bold=True, size=sizes.get(level, 12),
                 sb=befores.get(level, 6), sa=4)


def _shade(row_or_cell, fill="D9E2F3"):
    cells = row_or_cell.cells if hasattr(row_or_cell, "cells") else [row_or_cell]
    for cell in cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  fill)
        tcPr.append(shd)


def _table(doc, headers, rows, col_widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = table.rows[0]
    _shade(hrow, "1F3864")
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.width = Cm(col_widths[i])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        _font(r, bold=True, size=9, color=(255, 255, 255))

    for ri, row_data in enumerate(rows):
        row = table.add_row()
        if ri % 2 == 0:
            _shade(row, "EBF0FA")
        is_total = str(row_data[0]).startswith("**")
        if is_total:
            _shade(row, "D5E4F7")
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Cm(col_widths[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            clean = str(val).replace("**", "")
            r = p.add_run(clean)
            _font(r, bold=is_total, size=9)
    return table


def _code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'EBF0FA')
        pPr.append(shd)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)


def _field_run(paragraph, field_code):
    """在段落中插入 Word 域代码（如 PAGE、NUMPAGES）"""
    r = paragraph.add_run()
    r.font.size = Pt(9)
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_code} '
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    r._r.append(begin)
    r._r.append(instr)
    r._r.append(end)
    return r


def _setup_header_footer(doc, doc_number, product_name):
    """添加页眉（产品名 | 文件编号）和页脚（第 X 页 / 共 Y 页）"""

    def _plain(p, txt):
        r = p.add_run(txt)
        r.font.size  = Pt(9)
        r.font.name  = '宋体'
        r.font.color.rgb = RGBColor(100, 100, 100)

    for section in doc.sections:
        # ── 页眉 ──────────────────────────────────────────────────────────
        header = section.header
        if not header.paragraphs:
            hp = header.add_paragraph()
        else:
            hp = header.paragraphs[0]
        hp.clear()  # 清空默认内容
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after  = Pt(4)

        # 左：产品名 + Tab + 右：文件编号
        pPr = hp._p.get_or_add_pPr()
        tabs_el = OxmlElement('w:tabs')
        tab_el  = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:pos'), '8640')   # ~15.24cm 文字区宽度（A4，2.8cm边距）
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

        r_left = hp.add_run(product_name + '\t')
        r_left.font.size  = Pt(9)
        r_left.font.name  = '宋体'
        r_left.font.color.rgb = RGBColor(100, 100, 100)

        r_right = hp.add_run(f'文件编号：{doc_number}')
        r_right.font.size  = Pt(9)
        r_right.font.name  = '宋体'
        r_right.font.color.rgb = RGBColor(100, 100, 100)

        # 页眉分割线
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '4472C4')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # ── 页脚 ──────────────────────────────────────────────────────────
        footer = section.footer
        if not footer.paragraphs:
            fp = footer.add_paragraph()
        else:
            fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        fp.paragraph_format.space_after  = Pt(0)

        _plain(fp, '第 ')
        _field_run(fp, 'PAGE')
        _plain(fp, ' 页  /  共 ')
        _field_run(fp, 'NUMPAGES')
        _plain(fp, ' 页')


# ─────────────────────────────────────────────────────────────────────────────
# Word 文档构建
# ─────────────────────────────────────────────────────────────────────────────
def build_doc(cfg, res):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    pname   = cfg.get("product_name", "直流列头柜")
    project = cfg.get("project_name", "（项目名称）")
    docno   = cfg.get("doc_number",   "XXXX-MTBF-001")
    voltage = cfg.get("voltage",      "DC")
    csize   = cfg.get("cabinet_size", "—")
    entry   = cfg.get("entry_style",  "—")
    inp     = cfg["input_config"]
    lp      = cfg["lightning_protection"]
    mon     = cfg["monitoring"]
    has_imd = res["has_imd"]
    req_h   = res["req_h"]
    pq_i    = res["pq_i"]
    # 动态 λ 上限（对应 MTBF 要求）
    lam_req = round(1e6 / req_h, 4)

    verdict     = "满足 ✓" if res["pass"] else "不满足 ✗"
    margin_str  = (f'+{res["margin_pct"]}%（高于要求 {res["margin_h"]:,} h）'
                   if res["pass"] else
                   f'{res["margin_pct"]}%（低于要求 {abs(res["margin_h"]):,} h）')

    # 输出规格描述
    out_desc = "，".join(
        f'{g["count"]}路 {g["rating"]}/{g["poles"]}P {g.get("type","熔断器")}'
        for g in res["out_groups"]
    )
    mon_desc = f'{mon.get("hmi_size","7寸")}触摸屏监控' + ('，直流分路绝缘监测' if has_imd else '')

    # 页眉/页脚
    _setup_header_footer(doc, docno, pname)

    # ── 封面 ────────────────────────────────────────────────────────────────
    _para(doc, '')
    _para(doc, '[投标单位名称]', bold=True, size=16,
          align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
    _para(doc, pname, bold=True, size=20,
          align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=4, color=(31, 56, 100))
    _para(doc, '平均无故障间隔时间（MTBF）计算书', bold=True, size=16,
          align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=20)

    ct = doc.add_table(rows=7, cols=2)
    ct.style = 'Table Grid'
    ct.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ('文件编号', docno),
        ('项目名称', project),
        ('产品型号', pname),
        ('编制',     ''),
        ('审核',     ''),
        ('批准',     ''),
        ('日期',     '年    月    日'),
    ]):
        cells = ct.rows[i].cells
        cells[0].width = Cm(3)
        cells[1].width = Cm(11)
        for cell, txt, bold in [(cells[0], k, True), (cells[1], v, False)]:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font(p.add_run(txt), bold=bold, size=10.5)
    doc.add_page_break()

    # ── 1 概述 ──────────────────────────────────────────────────────────────
    _heading(doc, '1  概述')
    _heading(doc, '1.1 编制目的', 2)
    _para(doc,
        f'本文件依据{project}技术要求，对所投标的{pname}产品进行平均无故障间隔时间'
        f'（MTBF）的定量可靠性预计计算，以证明产品满足"机柜MTBF≥{req_h:,}h"的技术指标要求，'
        f'并为用户提供可靠性数据依据。', size=10.5)

    _heading(doc, '1.2 适用范围', 2)
    _para(doc,
        f'• 产品类型：{pname}\n'
        f'• 主要配置：{inp["count"]}路{inp["rating"]}{inp.get("type","熔断器")}输入 + '
        f'{out_desc}输出 + 直流防雷（In={lp["rating_kA"]}kA）+ {mon_desc}\n'
        f'• 应用场景：数据中心/通信机房，室内固定安装', size=10.5)

    # ── 2 计算依据 ───────────────────────────────────────────────────────────
    _heading(doc, '2  计算依据与方法')
    _heading(doc, '2.1 引用标准', 2)
    _table(doc,
        ['标准编号', '标准名称'],
        [
            ['MIL-HDBK-217F Notice 2', '《电子设备可靠性预计手册》（美国军用标准，国际通用）'],
            ['GB/T 17626.5',           '《电磁兼容 试验和测量技术》'],
            ['GB/T 11463',             '《电子测量仪器可靠性试验》'],
            ['YD/T 1360',              '《通信设备的可靠性要求和考核方法》'],
            ['IEC 61709',              '《电子元器件可靠性参考条件》'],
        ],
        [4.5, 11.5])
    doc.add_paragraph()

    _heading(doc, '2.2 计算方法', 2)
    _para(doc,
        '采用 MIL-HDBK-217F 零件计数法（Parts Count Method）进行系统级MTBF预计。'
        '对有制造商规格书的成套子组件（触摸屏、监控模块、绝缘监测模块、导轨电源），'
        '直接采用制造商声明的MTBF折算为失效率后代入计算。', size=10.5)
    _para(doc, '基本计算公式：', bold=True, size=10.5)
    _code(doc, [
        'λ_sys = Σ_i ( N_i × λ_Gi × π_Qi )',
        'MTBF  = 1 / λ_sys',
        '',
        '  λ_sys ：系统总失效率（×10⁻⁶/h）',
        '  N_i   ：第 i 类元器件数量',
        '  λ_Gi  ：GF 环境、40°C 下通用失效率（MIL-HDBK-217F 查表）',
        '  π_Qi  ：质量系数（本项目取 πQ=' + f'{pq_i}，{cfg.get("component_grade","国产工业级")}）',
    ])

    # ── 3 技术规格 ───────────────────────────────────────────────────────────
    _heading(doc, '3  产品技术规格')
    spec_rows = [
        ['产品名称',     pname],
        ['外形尺寸',     csize],
        ['开门方式',     entry],
        ['额定输入电压', voltage],
        ['输入支路规格', f'{inp["count"]}路，每路{inp["rating"]}{inp.get("type","熔断器")}保护'],
        ['输出支路规格', out_desc],
        ['防雷保护',     f'直流防雷器，In={lp["rating_kA"]}kA（{lp.get("waveform","8/20μs")}）'],
        ['监控单元',     f'{mon.get("hmi_size","7寸")}工业触摸屏，主支路监控' +
                         ('，直流分路绝缘监测' if has_imd else '')],
        ['安装方式',     '落地安装，室内数据中心/机房环境'],
        ['器件质量标准', f'{cfg.get("component_grade","国产工业级")}品牌器件，关键器件通过 CCC/CE 认证'],
    ]
    _table(doc, ['技术参数', '规格描述'], spec_rows, [5.5, 10.5])
    doc.add_paragraph()

    # ── 4 计算条件 ───────────────────────────────────────────────────────────
    _heading(doc, '4  计算条件与假设')
    _heading(doc, '4.1 工作环境条件', 2)
    _table(doc,
        ['参数', '取值', '说明'],
        [
            ['环境温度', f'{cfg.get("env_temp_c",40)}°C', '按机房最高设计温度取值，偏保守'],
            ['相对湿度', '≤85%（非凝结）', '机房正常运行环境'],
            ['使用环境', '地面固定（GF）', 'MIL-HDBK-217F 环境类别，πE=1.0'],
            ['散热方式', '自然对流', '柜体开孔通风，无强制风冷'],
            ['维护方式', '人工定期维护', '不计维修时间，按不可修复模型计算'],
        ],
        [4, 4, 8])
    doc.add_paragraph()

    _heading(doc, '4.2 质量等级假设', 2)
    _table(doc,
        ['元器件类别', '质量因子 πQ', '依据'],
        [
            ['通用无源元件（电阻、电容等）',          f'{pq_i}', '工业级认证品，符合 GB/T 标准'],
            ['通用有源元件（IC、传感器等）',           f'{pq_i}', '工业温度范围，通过可靠性筛选'],
            ['熔断器/断路器（工业级）',               f'{pq_i}', '通过 CCC/CE 认证，额定降额使用'],
            ['连接器/端子排',                         f'{pq_i}', '工业级压接端子，通过阻燃认证'],
            ['继电器',                                f'{pq_i}', '工业级密封继电器'],
            ['LED 指示灯',                            f'{PQ_LED}', '工业级LED，寿命>50,000h'],
            ['铜排/汇流母排',                         f'{PQ_BUSBAR}', '纯导体，无电子失效机制'],
            ['成套子组件（触摸屏、监控板、电源等）',   '—', '采用制造商声明MTBF数据'],
        ],
        [6.5, 2.5, 7])
    doc.add_paragraph()

    _heading(doc, '4.3 应力条件', 2)
    _para(doc,
        '所有元器件均按额定电气应力的 50%（应力比 S=0.5）进行降额使用，'
        '与 MIL-HDBK-217F 中 40°C 地面固定环境下的通用失效率（λ_G）基准一致。'
        '本预计属于设计阶段零件计数法，计算精度约为 ±20%。', size=10.5)

    # ── 5 子系统划分 ─────────────────────────────────────────────────────────
    _heading(doc, '5  系统划分与元器件清单')
    _para(doc,
        f'依据产品功能，将{pname}划分为以下两个子系统：\n'
        '• 子系统一（配电子系统）：承担直流电力分配功能，包含熔断器/断路器、母排、防雷器、连接端子等。\n'
        '• 子系统二（监控子系统）：承担数据采集、显示、绝缘监测与告警功能，包含触摸屏、'
        '监控主控板、传感器、导轨电源等。', size=10.5)

    # ── 6 计算过程 ───────────────────────────────────────────────────────────
    _heading(doc, '6  MTBF 计算过程')
    _heading(doc, '6.1 子系统一：配电子系统失效率计算', 2)
    _table(doc,
        ['序号', '元器件类型', '数量\nNi', 'λ_G\n(×10⁻⁶/h)', '质量因子\nπQ',
         '单件失效率\n(×10⁻⁶/h)', '合计\n(×10⁻⁶/h)'],
        res["ss1_rows"],
        [0.8, 5.0, 0.9, 1.4, 1.3, 1.7, 1.5])
    _para(doc,
        '注：MIL-HDBK-217F 熔断器基本失效率 λ_G=0.010×10⁻⁶/h（Section 12，GF，πE=1.0）；'
        '断路器 λ_G=0.010×10⁻⁶/h（Section 11）；防雷器（MOV）λ_G=0.003×10⁻⁶/h（Section 26）；'
        '连接器 λ_G=0.0027×10⁻⁶/h（Section 15）；内部线缆按工程经验取 0.20×10⁻⁶/h。', size=9)
    doc.add_paragraph()

    _heading(doc, '6.2 子系统二：监控子系统失效率计算', 2)
    _table(doc,
        ['序号', '元器件/组件类型', '数量', '参考MTBF或λ_G', 'πQ', '合计\n(×10⁻⁶/h)'],
        res["ss2_rows"],
        [0.8, 5.4, 0.9, 3.8, 0.8, 1.9])
    _para(doc,
        '注：触摸屏取工业级HMI典型声明值（威纶通、昆仑通态等）；'
        + ('绝缘监测模块取工业级IMD典型声明值（MTBF≥600,000h）；' if has_imd else '')
        + '导轨电源取MEAN WELL MDR系列Telcordia预计值；主控板按同类工业控制板卡设计估算。', size=9)
    doc.add_paragraph()

    _heading(doc, '6.3 系统总失效率汇总与 MTBF 计算', 2)
    _table(doc,
        ['子系统', '失效率 (×10⁻⁶/h)', '占比'],
        [
            ['子系统一：配电子系统', f'{res["lambda1"]:.4f}', f'{res["pct_ss1"]}%'],
            ['子系统二：监控子系统', f'{res["lambda2"]:.4f}', f'{res["pct_ss2"]}%'],
            ['**系统总计 λ_总**',    f'**{res["lambda_total"]:.4f}**', '**100%**'],
        ],
        [7, 4, 5])
    doc.add_paragraph()

    _para(doc, 'MTBF 计算过程：', bold=True, size=10.5)
    _code(doc, [
        f'λ_total = λ1 + λ2',
        f'        = {res["lambda1"]:.4f} + {res["lambda2"]:.4f}',
        f'        = {res["lambda_total"]:.4f} × 10⁻⁶ /h',
        '',
        f'MTBF = 1 / λ_total',
        f'     = 1 / ({res["lambda_total"]:.4f} × 10⁻⁶)',
        f'     = {res["mtbf_h"]:,} h',
        f'     ≈ {res["mtbf_wan"]} 万小时  ≈  {res["mtbf_years"]} 年（按 8760 h/年计）',
    ])

    # ── 7 结论 ──────────────────────────────────────────────────────────────
    _heading(doc, '7  结论')
    _table(doc,
        ['参数', '计算结果', '技术要求', '判定'],
        [
            ['系统总失效率 λ',
             f'{res["lambda_total"]:.4f} × 10⁻⁶ /h',
             f'≤{lam_req:.4f} × 10⁻⁶ /h',
             verdict],
            ['系统 MTBF',
             f'{res["mtbf_h"]:,} h',
             f'≥{req_h:,} h',
             verdict],
            ['可靠性裕量', margin_str, '', ''],
        ],
        [4.5, 4.5, 4, 3])
    doc.add_paragraph()

    _para(doc,
        f'{pname}   MTBF = {res["mtbf_h"]:,} h   【{verdict.replace(" ✓","").replace(" ✗","")} ≥{req_h:,} h 要求】',
        bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
    doc.add_paragraph()

    conclusion = (
        f'综上，本公司投标的{pname}，按MIL-HDBK-217F零件计数法预计的系统MTBF为'
        f'{res["mtbf_h"]:,}小时，{"满足" if res["pass"] else "未满足"}招标文件要求的'
        f'"机柜MTBF≥{req_h:,}h"指标。'
        + (f'可靠性裕量约为{res["margin_pct"]}%。' if res["pass"] else '')
    )
    _para(doc, conclusion, size=10.5)

    major_src = "监控子系统" if res["lambda2"] > res["lambda1"] else "配电子系统"
    major_pct = res["pct_ss2"] if res["lambda2"] > res["lambda1"] else res["pct_ss1"]
    _para(doc,
        f'影响系统可靠性的主要因素为{major_src}（占总失效率的{major_pct}%）。'
        '在实际工程中，可通过以下措施进一步提升可靠性：', size=10.5)
    suggestions = [
        '1. 选用MTBF≥500,000h的工业级触摸屏（如威纶通cMT系列）；',
        '2. 对监控供电电源进行N+1冗余配置；',
        '3. 对保护器件执行严格的额定降额（推荐应力比S≤0.5）。',
    ]
    if has_imd:
        suggestions.append('4. 选用MTBF≥1,000,000h的高可靠绝缘监测模块。')
    for item in suggestions:
        _para(doc, item, size=10.5)

    doc.add_page_break()

    # ── 附录A ───────────────────────────────────────────────────────────────
    _heading(doc, '附录A  主要参数引用说明')
    _heading(doc, 'A.1 MIL-HDBK-217F 元器件通用失效率（GF环境，40°C）', 2)
    _table(doc,
        ['元器件类型', 'MIL-HDBK-217F章节', 'λ_G (×10⁻⁶/h)', '备注'],
        [
            ['低压熔断器',        'Section 12', '0.0100', '地面固定，应力比0.5'],
            ['微型断路器（MCB）', 'Section 11', '0.0100', '2P/4P工业级，IEC 60947'],
            ['MOV/TVS 防雷器',    'Section 26', '0.0030', '浪涌抑制器'],
            ['铜排/导体',         'Section 17', '0.0010', '工程经验估算'],
            ['接插件/端子',       'Section 15', '0.0027', '通用矩形连接器'],
            ['霍尔传感器',        'Section 5',  '0.0100', '线性集成电路'],
            ['小型继电器',        'Section 13', '0.0400', '通用目的，密封型'],
            ['LED 指示灯',        'Section 6',  '0.0100', '高可靠LED'],
        ],
        [4.5, 3.5, 2.5, 5.5])
    doc.add_paragraph()

    _heading(doc, 'A.2 成套组件 MTBF 参考来源', 2)
    a2_rows = [
        ['7寸工业触摸屏',  '300,000 h',   '工业级HMI典型声明值；采购时以供应商数据为准'],
        ['10寸工业触摸屏', '250,000 h',   '工业级HMI典型声明值（大屏背光寿命略低）'],
        ['监控主控板',      '800,000 h',   '依据MIL-HDBK-217F Sec.5/9/10 对典型工业控制板估算'],
        ['导轨电源（DIN Rail）', '2,000,000 h', 'MEAN WELL MDR-60系列Telcordia预计值（40°C）'],
    ]
    if has_imd:
        a2_rows.append(['直流分路绝缘监测模块', '600,000 h', '工业级DC IMD典型声明值；采购时以供应商数据为准'])
    _table(doc, ['组件', '参考MTBF', '说明'], a2_rows, [4, 2.5, 9.5])
    doc.add_paragraph()

    _heading(doc, 'A.3 计算假设说明', 2)
    for item in [
        f'1. 本计算基于零件计数法，属于设计阶段的可靠性预计，计算精度约 ±20%，实际MTBF受具体元器件品牌、工艺影响，可能有所差异。',
        f'2. 计算中所有元器件均假设工作于稳定运行状态（不计开关机应力），符合指数分布失效规律。',
        f'3. 质量因子πQ={pq_i}对应{cfg.get("component_grade","国产工业级")}元器件，适用于通过CCC/CE/GB认证的工业品。',
        f'4. 中标后，投标方将根据实际采购BOM及元器件规格书，提交更精确的MTBF详细计算书（含每一元器件明细）作为最终交付文件。',
    ]:
        _para(doc, item, size=10.5)

    doc.add_paragraph()
    _para(doc, '以上计算结果由本公司可靠性工程师核算，数据真实、方法规范，特此确认。', size=10.5)
    doc.add_paragraph()

    sig = doc.add_table(rows=2, cols=3)
    sig.style = 'Table Grid'
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, txt in zip(sig.rows[0].cells, ['编制（签名）', '审核（签名）', '批准（签名/公章）']):
        cell.width = Cm(5.3)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(p.add_run(txt), bold=True, size=10.5)
    for cell in sig.rows[1].cells:
        cell.width = Cm(5.3)
        cell.paragraphs[0].add_run('\n')

    doc.add_paragraph()
    _para(doc, '日期：______年______月______日', size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# 公开接口
# ─────────────────────────────────────────────────────────────────────────────
def run(cfg):
    """计算 + 生成文档，返回 (output_path, result_dict)"""
    res     = calculate(cfg)
    doc     = build_doc(cfg, res)
    outpath = cfg.get("output_path", os.path.join(os.getcwd(), "MTBF计算书.docx"))
    doc.save(outpath)
    return outpath, res


# ─────────────────────────────────────────────────────────────────────────────
# CLI 入口
# ─────────────────────────────────────────────────────────────────────────────
DEFAULT_CONFIG = {
    "project_name":         "（填写项目名称）",
    "product_name":         "240V直流列头柜",
    "doc_number":           "XXXX-MTBF-001",
    "output_path":          "D:/MTBF计算书.docx",
    "voltage":              "DC 240V",
    "cabinet_size":         "1200×300×2200mm",
    "entry_style":          "前开门，上进线上出线",
    "mtbf_requirement_h":   100_000,        # MTBF 指标要求（小时）
    "component_grade":      "国产工业级",    # 影响 πQ：国产工业级=2.0，进口工业级=1.5，军品级=1.0
    "env_temp_c":           40,
    "input_config": {
        "count":  2,
        "rating": "500A",
        "type":   "双极熔断器"
    },
    # 单组输出：
    "output_config": {
        "total_count":        60,
        "rating_per_circuit": "63A",
        "poles":              2,
        "type":               "直流开关"
    },
    # 混合输出示例（取消注释并删除上方 output_config 即可使用）：
    # "output_config": {
    #     "groups": [
    #         {"count": 36, "rating": "125A", "poles": 1, "type": "熔断器"},
    #         {"count": 36, "rating": "63A",  "poles": 1, "type": "熔断器"}
    #     ]
    # },
    "lightning_protection": {
        "count":     2,
        "rating_kA": 20,
        "waveform":  "8/20μs"
    },
    "monitoring": {
        "hmi_size":               "7寸",
        "has_insulation_monitor": True,
        "hall_sensor_count":      4,    # 默认 = n_in × 2（正负母线各一）
        "relay_count":            4,
        "led_count":              10
    }
}

if __name__ == "__main__":
    if len(sys.argv) > 1:
        try:
            cfg = json.loads(sys.argv[1])
        except json.JSONDecodeError as e:
            sys.stderr.write(f"[ERROR] JSON 解析失败: {e}\n")
            sys.exit(1)
    else:
        cfg = DEFAULT_CONFIG
    outpath, res = run(cfg)
    out = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    out.write(f"文档已生成: {outpath}\n")
    out.write(f"MTBF = {res['mtbf_h']:,} h  ({res['mtbf_wan']} 万小时)\n")
    out.write(f"lambda_total = {res['lambda_total']:.4f} x10^-6/h\n")
    out.write(f"判定: {'满足' if res['pass'] else '不满足'} >= {res['req_h']:,} h\n")
    out.flush()
