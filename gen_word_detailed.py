#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('技术规范提取报告 - 超详细版')
run.font.size = Pt(22)
run.font.bold = True

# 项目信息
doc.add_heading('项目信息', level=1)
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Light Grid Accent 1'
info_data = [
    ('项目名称', '2025年汕头联通通信楼机电工程建设项目EPC总承包项目'),
    ('技术规范书', '发包人要求(技术规范书)-20251003.doc'),
    ('提取日期', '2025-03-17'),
    ('规范页数', '272页'),
    ('文档字数', '约30,813字')
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v

# 第一章 低压开关柜
doc.add_heading('第一章 低压开关柜（6.3）', level=1)

doc.add_heading('1.1 技术标准', level=2)
doc.add_paragraph('应符合 GB 50054《低压配电设计规范》、GB 50045《高层民用建筑设计防火规范》等国家标准。')

doc.add_heading('1.2 电气性能参数', level=2)
table1 = doc.add_table(rows=10, cols=3)
table1.style = 'Light Grid Accent 1'
table1.rows[0].cells[0].text = '参数'
table1.rows[0].cells[1].text = '规范要求'
table1.rows[0].cells[2].text = '原文参考'

data1 = [
    ('额定电压', '≥380V', '额定电压≥380V'),
    ('绝缘电压', '≥1000V', '额定电压≥380V'),
    ('短时耐受电流', '≥31.5kA(4S)', '短时耐受电流≥31.5kA(4S)'),
    ('峰值耐受电流', '≥80kA', '峰值耐受电流≥80kA'),
    ('框架断路器', '≥800A', '框架断路器≥800A'),
    ('防护等级', '≥IP30', '防护等级不小于IP30'),
    ('频率', '50Hz', '50Hz'),
    ('系统', '三相五线制', '三相五线制'),
    ('绝缘电压', '≥1000V', '额定电压≥380V'),
]
for i, row in enumerate(data1, 1):
    for j, val in enumerate(row):
        table1.rows[i].cells[j].text = val

doc.add_heading('1.3 导体材料要求', level=2)
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = '参数'
table2.rows[0].cells[1].text = '规范要求'
table2.rows[0].cells[2].text = '原文参考'
data2 = [
    ('导体材质', 'T2电解铜', 'T2电解铜，含铜率不低于99.95%'),
    ('纯度', '≥99.95%', '含铜率不低于99.95%'),
    ('导电率', '≥97%IACS', 'T2电解铜'),
]
for i, row in enumerate(data2, 1):
    for j, val in enumerate(row):
        table2.rows[i].cells[j].text = val

doc.add_heading('1.4 SVG/APF混合滤波补偿柜', level=2)
table3 = doc.add_table(rows=8, cols=2)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = '参数'
table3.rows[0].cells[1].text = '规范要求'
data3 = [
    ('电压范围', '380V±20%'),
    ('频率范围', '50Hz±10%'),
    ('全响应时间', '<10ms'),
    ('功率损耗', '<3%设备额定容量'),
    ('输出电流不平衡度', '100%载时小于3%'),
    ('APF谐波补偿范围', '2~50次'),
    ('SVG谐波补偿范围', '2~25次'),
]
for i, row in enumerate(data3, 1):
    for j, val in enumerate(row):
        table3.rows[i].cells[j].text = val

# 第二章 配电小母线
doc.add_heading('第二章 配电小母线（6.4）', level=1)

doc.add_heading('2.1 环境条件', level=2)
table4 = doc.add_table(rows=5, cols=2)
table4.style = 'Light Grid Accent 1'
table4.rows[0].cells[0].text = '参数'
table4.rows[0].cells[1].text = '规范要求'
data4 = [
    ('工作温度', '-10℃～45℃'),
    ('相对湿度', '≤85%RH（25℃±5℃时）'),
    ('储运温度', '-40℃～+70℃'),
    ('海拔高度', '≤1000m'),
]
for i, row in enumerate(data4, 1):
    for j, val in enumerate(row):
        table4.rows[i].cells[j].text = val

doc.add_heading('2.2 工作电源', level=2)
table5 = doc.add_table(rows=4, cols=2)
table5.style = 'Light Grid Accent 1'
table5.rows[0].cells[0].text = '参数'
table5.rows[0].cells[1].text = '规范要求'
data5 = [
    ('单相电压', '187V-242V'),
    ('三相电压', '323V-418V'),
    ('频率', '50Hz'),
]
for i, row in enumerate(data5, 1):
    for j, val in enumerate(row):
        table5.rows[i].cells[j].text = val

doc.add_heading('2.3 母线槽技术要求', level=2)
table6 = doc.add_table(rows=8, cols=2)
table6.style = 'Light Grid Accent 1'
table6.rows[0].cells[0].text = '参数'
table6.rows[0].cells[1].text = '规范要求'
data6 = [
    ('导体材质', 'T2电解铜'),
    ('含铜率', '≥99.95%'),
    ('线制', '3L+N+PE'),
    ('N线容量', '100%相线容量'),
    ('PE线容量', '50%相线容量'),
    ('绝缘等级', 'B级及以上'),
    ('老化寿命', '≥25年'),
]
for i, row in enumerate(data6, 1):
    for j, val in enumerate(row):
        table6.rows[i].cells[j].text = val

# 第三章 模块化UPS
doc.add_heading('第三章 模块化UPS电源（6.5）', level=1)

doc.add_heading('3.1 输入特性', level=2)
table7 = doc.add_table(rows=4, cols=3)
table7.style = 'Light Grid Accent 1'
table7.rows[0].cells[0].text = '参数'
table7.rows[0].cells[1].text = '规范要求'
table7.rows[0].cells[2].text = '备注'
data7 = [
    ('输入电压范围', '-15%～+10%', '输出额定容量'),
    ('输入频率范围', '50Hz±2.5Hz', '-'),
    ('输入功率因数', '见下表', '100%/75%/50%负载'),
]
for i, row in enumerate(data7, 1):
    for j, val in enumerate(row):
        table7.rows[i].cells[j].text = val

doc.add_heading('输入功率因数表', level=3)
table8 = doc.add_table(rows=4, cols=2)
table8.style = 'Light Grid Accent 1'
table8.rows[0].cells[0].text = '负载百分比'
table8.rows[0].cells[1].text = '功率因数要求'
data8 = [
    ('100%', '≥0.99'),
    ('75%', '≥0.97'),
    ('50%', '≥0.94'),
]
for i, row in enumerate(data8, 1):
    for j, val in enumerate(row):
        table8.rows[i].cells[j].text = val

doc.add_heading('3.2 输出特性', level=2)
table9 = doc.add_table(rows=4, cols=2)
table9.style = 'Light Grid Accent 1'
table9.rows[0].cells[0].text = '参数'
table9.rows[0].cells[1].text = '规范要求'
data9 = [
    ('稳压精度', '±1%'),
    ('输出电压范围', '-43.2V～-57.6V'),
    ('输出功率因数', '≥0.90'),
]
for i, row in enumerate(data9, 1):
    for j, val in enumerate(row):
        table9.rows[i].cells[j].text = val

doc.add_heading('3.3 系统效率', level=2)
table10 = doc.add_table(rows=5, cols=2)
table10.style = 'Light Grid Accent 1'
table10.rows[0].cells[0].text = '负载百分比'
table10.rows[0].cells[1].text = '效率要求'
data10 = [
    ('100%负载', '≥94.5%'),
    ('50%负载', '≥95%'),
    ('30%负载', '≥94%'),
    ('20%负载', '≥93%'),
]
for i, row in enumerate(data10, 1):
    for j, val in enumerate(row):
        table10.rows[i].cells[j].text = val

doc.add_heading('3.4 整流模块效率', level=2)
table11 = doc.add_table(rows=5, cols=2)
table11.style = 'Light Grid Accent 1'
table11.rows[0].cells[0].text = '负载百分比'
table11.rows[0].cells[1].text = '效率要求'
data11 = [
    ('100%负载', '≥95.5%'),
    ('50%负载', '≥96%'),
    ('30%负载', '≥95%'),
    ('20%负载', '≥94%'),
]
for i, row in enumerate(data11, 1):
    for j, val in enumerate(row):
        table11.rows[i].cells[j].text = val

# 第四章 240V直流列头柜
doc.add_heading('第四章 240V直流列头柜（6.6）', level=1)

doc.add_heading('4.1 关键技术指标', level=2)
table12 = doc.add_table(rows=7, cols=2)
table12.style = 'Light Grid Accent 1'
table12.rows[0].cells[0].text = '参数'
table12.rows[0].cells[1].text = '规范要求'
data12 = [
    ('工作电压', '≥240V DC'),
    ('极限短路分断能力', '≥10kA'),
    ('防雷器件', 'In=20kA (8/20μs)'),
    ('电压降', '≤0.3V (20℃,额定电流)'),
    ('MTBF', '≥100,000h'),
    ('通信接口', 'RS485 Modbus'),
]
for i, row in enumerate(data12, 1):
    for j, val in enumerate(row):
        table12.rows[i].cells[j].text = val

# 第五章 48V直流系统
doc.add_heading('第五章 48V直流系统（6.7）', level=1)

doc.add_heading('5.1 系统参数', level=2)
table13 = doc.add_table(rows=3, cols=2)
table13.style = 'Light Grid Accent 1'
table13.rows[0].cells[0].text = '参数'
table13.rows[0].cells[1].text = '规范要求'
data13 = [
    ('工作电压范围', '-43.2V～-57.6V'),
    ('直流电压输出', '可手动或自动连续可调'),
]
for i, row in enumerate(data13, 1):
    for j, val in enumerate(row):
        table13.rows[i].cells[j].text = val

doc.add_heading('5.2 系统效率', level=2)
table14 = doc.add_table(rows=5, cols=2)
table14.style = 'Light Grid Accent 1'
table14.rows[0].cells[0].text = '负载百分比'
table14.rows[0].cells[1].text = '效率要求'
data14 = [
    ('100%负载', '≥94.5%'),
    ('50%负载', '≥95%'),
    ('30%负载', '≥94%'),
    ('20%负载', '≥93%'),
]
for i, row in enumerate(data14, 1):
    for j, val in enumerate(row):
        table14.rows[i].cells[j].text = val

doc.add_heading('5.3 整流模块效率', level=2)
table15 = doc.add_table(rows=5, cols=2)
table15.style = 'Light Grid Accent 1'
table15.rows[0].cells[0].text = '负载百分比'
table15.rows[0].cells[1].text = '效率要求'
data15 = [
    ('100%负载', '≥95.5%'),
    ('50%负载', '≥96%'),
    ('30%负载', '≥95%'),
    ('20%负载', '≥94%'),
]
for i, row in enumerate(data15, 1):
    for j, val in enumerate(row):
        table15.rows[i].cells[j].text = val

doc.add_heading('5.4 功率因数', level=2)
table16 = doc.add_table(rows=3, cols=2)
table16.style = 'Light Grid Accent 1'
table16.rows[0].cells[0].text = '参数'
table16.rows[0].cells[1].text = '规范要求'
data16 = [
    ('输入功率因数', '≥0.99'),
    ('谐波', '≤5%'),
]
for i, row in enumerate(data16, 1):
    for j, val in enumerate(row):
        table16.rows[i].cells[j].text = val

doc.add_heading('5.5 整流架容量', level=2)
table17 = doc.add_table(rows=4, cols=2)
table17.style = 'Light Grid Accent 1'
table17.rows[0].cells[0].text = '参数'
table17.rows[0].cells[1].text = '规范要求'
data17 = [
    ('整流架容量', '≥2000A'),
    ('交流配电', '380V/630A'),
    ('直流配电', '3000A'),
]
for i, row in enumerate(data17, 1):
    for j, val in enumerate(row):
        table17.rows[i].cells[j].text = val

# 第六章 48V直流列头柜
doc.add_heading('第六章 48V直流列头柜（6.8）', level=1)

doc.add_heading('6.1 关键技术指标', level=2)
table18 = doc.add_table(rows=6, cols=2)
table18.style = 'Light Grid Accent 1'
table18.rows[0].cells[0].text = '参数'
table18.rows[0].cells[1].text = '规范要求'
data18 = [
    ('工作电压', '-48V DC (-40V～-58V)'),
    ('电压降', '≤0.30V (20℃,额定电流)'),
    ('绝缘电阻', '≥30MΩ（对地）'),
    ('耐压值', '≥1000V'),
    ('防护等级', 'IP20'),
]
for i, row in enumerate(data18, 1):
    for j, val in enumerate(row):
        table18.rows[i].cells[j].text = val

# 第七章 配电箱柜
doc.add_heading('第七章 配电箱（柜）系统（6.11）', level=1)

doc.add_heading('7.1 基本参数', level=2)
table19 = doc.add_table(rows=4, cols=2)
table19.style = 'Light Grid Accent 1'
table19.rows[0].cells[0].text = '参数'
table19.rows[0].cells[1].text = '规范要求'
data19 = [
    ('额定电压', '400V/660V'),
    ('频率', '50Hz'),
    ('系统', '三相五线制'),
]
for i, row in enumerate(data19, 1):
    for j, val in enumerate(row):
        table19.rows[i].cells[j].text = val

doc.add_heading('7.2 断路器要求', level=2)
table20 = doc.add_table(rows=3, cols=2)
table20.style = 'Light Grid Accent 1'
table20.rows[0].cells[0].text = '参数'
table20.rows[0].cells[1].text = '规范要求'
data20 = [
    ('框架断路器Ics', '≥50kA'),
    ('塑壳断路器Ics', '≥36kA'),
]
for i, row in enumerate(data20, 1):
    for j, val in enumerate(row):
        table20.rows[i].cells[j].text = val

doc.add_heading('7.3 防护等级', level=2)
table21 = doc.add_table(rows=3, cols=2)
table21.style = 'Light Grid Accent 1'
table21.rows[0].cells[0].text = '环境'
table21.rows[0].cells[1].text = '防护等级要求'
data21 = [
    ('潮湿环境', '≥IP55'),
    ('其他环境', '≥IP40'),
]
for i, row in enumerate(data21, 1):
    for j, val in enumerate(row):
        table21.rows[i].cells[j].text = val

doc.add_heading('7.4 抗震要求', level=2)
table22 = doc.add_table(rows=2, cols=2)
table22.style = 'Light Grid Accent 1'
table22.rows[0].cells[0].text = '参数'
table22.rows[0].cells[1].text = '规范要求'
table22.rows[1].cells[0].text = '抗震等级'
table22.rows[1].cells[1].text = '8级'

# 第八章 PDU
doc.add_heading('第八章 PDU（6.27.3）', level=1)

doc.add_heading('8.1 输入要求', level=2)
table23 = doc.add_table(rows=4, cols=2)
table23.style = 'Light Grid Accent 1'
table23.rows[0].cells[0].text = '参数'
table23.rows[0].cells[1].text = '规范要求'
data23 = [
    ('输入电压', '220V AC'),
    ('输入电流', '63A'),
    ('频率', '50-60Hz'),
]
for i, row in enumerate(data23, 1):
    for j, val in enumerate(row):
        table23.rows[i].cells[j].text = val

doc.add_heading('8.2 输出配置', level=2)
table24 = doc.add_table(rows=5, cols=2)
table24.style = 'Light Grid Accent 1'
table24.rows[0].cells[0].text = '参数'
table24.rows[0].cells[1].text = '规范要求'
data24 = [
    ('10A输出孔位', '16位'),
    ('16A输出孔位', '8位'),
    ('空开品牌', 'ABB/施耐德/西门子'),
    ('阻燃等级', 'V0级'),
]
for i, row in enumerate(data24, 1):
    for j, val in enumerate(row):
        table24.rows[i].cells[j].text = val

# 综合风险分析
doc.add_heading('综合风险分析', level=1)

doc.add_heading('风险等级汇总', level=2)
table25 = doc.add_table(rows=9, cols=3)
table25.style = 'Light Grid Accent 1'
table25.rows[0].cells[0].text = '章节'
table25.rows[0].cells[1].text = '风险等级'
table25.rows[0].cells[2].text = '主要问题'
risk_data = [
    ('6.3 低压开关柜', 'MEDIUM', 'SVG/APF混合滤波需特殊配置'),
    ('6.4 配电小母线', 'LOW', '标准配置'),
    ('6.5 UPS', 'MEDIUM', '效率指标需验证'),
    ('6.6 240V直流', 'LOW', '标准配置'),
    ('6.7 48V直流系统', 'MEDIUM', '系统效率和模块效率需验证'),
    ('6.8 48V直流列头柜', 'LOW', '标准配置'),
    ('6.11 配电箱', 'LOW', '标准配置'),
    ('6.27.3 PDU', 'LOW', '标准配置'),
]
for i, row in enumerate(risk_data, 1):
    for j, val in enumerate(row):
        table25.rows[i].cells[j].text = val

doc.add_heading('总体风险等级：MEDIUM', level=2)

doc.add_heading('建议', level=2)
doc.add_paragraph('1. 报价建议：按标准配置报价，无特殊成本增加', style='List Bullet')
doc.add_paragraph('2. 采购建议：常规供应商可满足，无需特殊定制', style='List Bullet')
doc.add_paragraph('3. 厂验重点：', style='List Bullet')
doc.add_paragraph('   - UPS效率指标（94.5%/95%/94%/93%）', style='List Bullet')
doc.add_paragraph('   - 48V系统效率指标', style='List Bullet')
doc.add_paragraph('   - 整流模块效率指标（95.5%/96%/95%/94%）', style='List Bullet')
doc.add_paragraph('   - SVG/APF混合滤波性能测试', style='List Bullet')

# 页脚
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('报告生成时间：2025-03-17  |  报告版本：3.0（超详细版）')

doc.save('extract_spec_report_detailed.docx')
print('Word文档已生成：extract_spec_report_detailed.docx')
